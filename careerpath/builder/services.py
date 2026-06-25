from io import BytesIO

from django.utils.text import slugify

from .models import ResumeProfile


def get_or_create_resume_profile(user):
    latest_result = user.assessment_results.first()
    defaults = {
        "headline": latest_result.career_name if latest_result else "",
        "summary": (
            f"Aspiring {latest_result.career_name} focused on steady skill growth and consistent execution."
            if latest_result
            else ""
        ),
        "portfolio_url": "",
    }
    profile, _ = ResumeProfile.objects.get_or_create(user=user, defaults=defaults)
    return profile


def resume_filename(profile, extension):
    user_name = profile.user.get_full_name() or profile.user.username
    base = slugify(f"{user_name} resume") or "mentoraa-resume"
    return f"{base}.{extension}"


def build_resume_docx(profile):
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10.5)

    user_name = profile.user.get_full_name() or profile.user.username
    document.add_heading(user_name, level=0)
    if profile.headline:
        document.add_paragraph(profile.headline)

    contact_parts = [
        profile.location,
        profile.phone,
        profile.user.email,
        profile.linkedin_url,
        profile.github_url,
        profile.portfolio_url,
    ]
    document.add_paragraph(" | ".join(part for part in contact_parts if part))

    if profile.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(profile.summary)

    education_items = profile.education_items.all()
    if education_items:
        document.add_heading("Education", level=1)
        for item in education_items:
            paragraph = document.add_paragraph()
            paragraph.add_run(item.institution).bold = True
            paragraph.add_run(f" - {item.degree}")
            years = f"{item.start_year}"
            if item.end_year:
                years += f" - {item.end_year}"
            document.add_paragraph(years)
            if item.description:
                document.add_paragraph(item.description, style="List Bullet")

    experience_items = profile.experience_items.all()
    if experience_items:
        document.add_heading("Experience", level=1)
        for item in experience_items:
            paragraph = document.add_paragraph()
            paragraph.add_run(item.title).bold = True
            paragraph.add_run(f" - {item.organization}")
            dates = item.start_date.strftime("%b %Y")
            if item.end_date:
                dates += f" - {item.end_date.strftime('%b %Y')}"
            document.add_paragraph(dates)
            for bullet in _split_resume_bullets(item.description):
                document.add_paragraph(bullet, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_resume_pdf(profile):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=42,
        leftMargin=42,
        topMargin=42,
        bottomMargin=42,
    )
    styles = getSampleStyleSheet()
    story = []

    user_name = profile.user.get_full_name() or profile.user.username
    story.append(Paragraph(user_name, styles["Title"]))
    if profile.headline:
        story.append(Paragraph(profile.headline, styles["Heading2"]))

    contact_parts = [
        profile.location,
        profile.phone,
        profile.user.email,
        profile.linkedin_url,
        profile.github_url,
        profile.portfolio_url,
    ]
    story.append(Paragraph(" | ".join(part for part in contact_parts if part), styles["Normal"]))
    story.append(Spacer(1, 12))

    if profile.summary:
        story.append(Paragraph("Summary", styles["Heading2"]))
        story.append(Paragraph(profile.summary, styles["Normal"]))
        story.append(Spacer(1, 10))

    education_items = profile.education_items.all()
    if education_items:
        story.append(Paragraph("Education", styles["Heading2"]))
        for item in education_items:
            years = f"{item.start_year}" + (f" - {item.end_year}" if item.end_year else "")
            story.append(Paragraph(f"<b>{item.institution}</b> - {item.degree}", styles["Normal"]))
            story.append(Paragraph(years, styles["Normal"]))
            if item.description:
                story.append(Paragraph(item.description, styles["Normal"]))
            story.append(Spacer(1, 8))

    experience_items = profile.experience_items.all()
    if experience_items:
        story.append(Paragraph("Experience", styles["Heading2"]))
        for item in experience_items:
            dates = item.start_date.strftime("%b %Y")
            if item.end_date:
                dates += f" - {item.end_date.strftime('%b %Y')}"
            story.append(Paragraph(f"<b>{item.title}</b> - {item.organization}", styles["Normal"]))
            story.append(Paragraph(dates, styles["Normal"]))
            for bullet in _split_resume_bullets(item.description):
                story.append(Paragraph(f"- {bullet}", styles["Normal"]))
            story.append(Spacer(1, 8))

    document.build(story)
    buffer.seek(0)
    return buffer


def _split_resume_bullets(text):
    if not text:
        return []
    return [line.strip(" -\t") for line in text.splitlines() if line.strip(" -\t")]
